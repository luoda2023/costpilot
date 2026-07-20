"""
造价通 - 8 类格式谱 Markdown 渲染器

把模板的 content_md 中的 Markdown 标记转为 docx 段落/表格/列表;

支持:
  - ##/###/####  章节标题
  - 段落 / 列表(- 与 1. )
  - **粗体** / *斜体*
  - 表格( | a | b | )
  - 代码块 / 引用块
  - 字段占位 {{key}} 替换为字段值
"""
import re
from pathlib import Path
from typing import Dict, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from packages.server.db.database import SessionLocal
from packages.server.db.models import Template, TemplateField


# ---------------------------------------------------------------------------
# Markdown 简易解析
# ---------------------------------------------------------------------------

def _apply_inline_formatting(paragraph, text):
    """处理 **粗体**、*斜体*、`代码`"""
    # 切分:粗体 / 斜体 / 代码
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(tok)


def _add_md_table(doc, header_row, data_rows):
    """添加 markdown 表格 -> docx 表格"""
    headers = [c.strip() for c in header_row.strip("|").split("|")]
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 数据
    for ri, row in enumerate(data_rows, start=1):
        cells = [c.strip() for c in row.strip("|").split("|")]
        for ci, v in enumerate(cells[:len(headers)]):
            table.rows[ri].cells[ci].text = v
    return table


def render_markdown_to_docx(doc: Document, md_text: str):
    """把 markdown 文本逐行写入 doc"""
    lines = md_text.split("\n")
    i = 0
    pending_table_header = None
    pending_table_rows = []

    def flush_table():
        nonlocal pending_table_header, pending_table_rows
        if pending_table_header:
            _add_md_table(doc, pending_table_header, pending_table_rows)
            pending_table_header = None
            pending_table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 表格检测
        if stripped.startswith("|") and stripped.endswith("|"):
            if pending_table_header is None:
                pending_table_header = stripped
            else:
                pending_table_rows.append(stripped)
            i += 1
            continue
        else:
            flush_table()

        # 空行
        if not stripped:
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_heading(level=level)
            run = p.add_run(text)
            # 中文字体
            from docx.oxml.ns import qn
            run.font.name = "宋体"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.left_indent = Cm(1)
            _apply_inline_formatting(p, stripped[1:].strip())
            i += 1
            continue

        # 无序列表
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline_formatting(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # 有序列表
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _apply_inline_formatting(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # 代码块
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _apply_inline_formatting(p, stripped)
        i += 1

    flush_table()


# ---------------------------------------------------------------------------
# 字段值替换
# ---------------------------------------------------------------------------

def fill_placeholders(md_text: str, fields: Dict[str, str]) -> str:
    """把 {{key}} 占位替换为字段值;未填的字段用 ____ 替代"""
    def replace(m):
        key = m.group(1).strip()
        return str(fields.get(key, f"__{key}__"))
    return re.sub(r"\{\{\s*([^}]+?)\s*\}\}", replace, md_text)


# ---------------------------------------------------------------------------
# 主导出函数
# ---------------------------------------------------------------------------

def export_template_to_docx(
    template_id: int,
    field_values: Dict[str, str],
    output_path: Optional[str] = None,
) -> str:
    """根据模板 ID + 字段值,生成 docx 文档"""
    session = SessionLocal()
    try:
        tpl = session.query(Template).get(template_id)
        if not tpl:
            raise ValueError(f"模板 ID={template_id} 不存在")

        # 字段定义(供默认值 / 必填校验)
        field_defs = session.query(TemplateField).filter_by(template_id=template_id).all()

        # 合并默认值
        merged = {f.field_key: f.default_value or "" for f in field_defs}
        merged.update(field_values or {})

        # 替换占位
        md_filled = fill_placeholders(tpl.content_md or "", merged)

        # 应用名作文件名
        proj_name = merged.get("project_name", "未命名")
        type_name = tpl.name

        # 生成文档
        doc = Document()
        # 默认字体
        from docx.oxml.ns import qn
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(11)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 封面
        doc.add_paragraph()
        doc.add_paragraph()
        title_p = doc.add_heading(level=0)
        title_run = title_p.add_run(f"{proj_name}\n{type_name}")
        title_run.font.size = Pt(28)
        title_run.font.name = "宋体"
        title_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元数据
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p.add_run(
            f"\n\n编制地区: {merged.get('region','')}  |  "
            f"编制阶段: {merged.get('stage','')}  |  "
            f"编制人: {merged.get('compiled_by','')}  |  "
            f"编制日期: {merged.get('compiled_at','')}"
        )

        doc.add_page_break()

        # 主体内容
        render_markdown_to_docx(doc, md_filled)

        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"./{proj_name}_{type_name}_{timestamp}.docx"

        doc.save(output_path)
        return str(Path(output_path).resolve())

    finally:
        session.close()
