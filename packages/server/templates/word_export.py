"""
造价通 - Word 报价表导出

输出结构:
  1) 封面(项目名 + 编制信息)
  2) 编制说明
  3) 投标报价汇总表
  4) 分部分项工程量清单计价表
  5) 措施项目清单计价表
  6) 规费/税金项目清单计价表

依赖: python-docx
"""
from pathlib import Path
from typing import Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

from packages.server.templates.pricing import QuoteResult


# ---------------------------------------------------------------------------
# 文档样式
# ---------------------------------------------------------------------------

def _set_font(run, name="宋体", size=11, bold=False):
    """中文字体同时设 eastAsia 和 ascii"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def _add_title(doc, text, size=18, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    _set_font(run, size=size, bold=True)
    return p


def _add_para(doc, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Pt(indent)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    return p


def _add_table(doc, headers, rows, col_widths_cm=None):
    """添加带表头样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 表头
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        _set_font(run, size=10, bold=True)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run("" if val is None else str(val))
            _set_font(run, size=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 列宽
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for r in table.rows:
                r.cells[i].width = Cm(w)

    return table


def _format_amount(x: float) -> str:
    return f"{x:,.2f}"


# ---------------------------------------------------------------------------
# 主导出函数
# ---------------------------------------------------------------------------

def export_quote_to_word(
    quote: QuoteResult,
    project_info: dict,
    output_path: Optional[str] = None,
) -> str:
    """导出完整报价表到 Word .docx"""
    if output_path is None:
        output_path = f"./{project_info.get('name','未命名')}_报价表_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    doc = Document()

    # 默认正文字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ------- 封面 -------
    doc.add_paragraph()
    doc.add_paragraph()
    _add_title(doc, "工程造价报价表", size=24)
    doc.add_paragraph()
    doc.add_paragraph()

    info_rows = [
        ("项目名称:", project_info.get("name", "")),
        ("项目所在地:", project_info.get("region", "")),
        ("工程阶段:", project_info.get("stage", "")),
        ("计税方式:", project_info.get("tax_method", "一般计税")),
        ("编制人:", project_info.get("compiler", "")),
        ("审核人:", project_info.get("reviewer", "")),
        ("编制日期:", project_info.get("date", datetime.now().strftime("%Y-%m-%d"))),
    ]
    info_table = _add_table(doc, ["项", "内容"], info_rows, col_widths_cm=[4, 12])
    doc.add_page_break()

    # ------- 编制说明 -------
    _add_title(doc, "编制说明", size=16, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()
    notes = [
        "一、本报价依据《建设工程工程量清单计价规范》GB50500-2013 编制。",
        f"二、项目所在地: {project_info.get('region', '')},工程阶段: {project_info.get('stage', '')}。",
        f"三、计税方式: {project_info.get('tax_method', '一般计税')}",
        "四、综合单价 = 材料费 + 人工费 + 机械费 + 管理费(5%) + 利润(5%) + 规费 + 税金(9%/3%)",
        "五、二类措施费以分部分项合计为基数,各项费率按地区取定;",
        "六、三类规费以分部分项+措施为基数,税金以分部分项+措施+规费为基数;",
        f"七、工程总造价: ¥ {_format_amount(quote.grand_total)} 元",
        f"    一类(分部分项): ¥ {_format_amount(quote.category1)}",
        f"    二类(措施费):    ¥ {_format_amount(quote.category2)}",
        f"    三类(规费+税金): ¥ {_format_amount(quote.category3)}",
    ]
    for note in notes:
        _add_para(doc, note, size=11, indent=22)
    doc.add_page_break()

    # ------- 投标报价汇总表 -------
    _add_title(doc, "投标报价汇总表")
    summary_rows = [
        ["1", "一类 分部分项工程费", _format_amount(quote.category1),
         f"{quote.category1 / quote.grand_total * 100:.2f}%" if quote.grand_total else "0%"],
        ["2", "二类 措施项目费", _format_amount(quote.category2),
         f"{quote.category2 / quote.grand_total * 100:.2f}%" if quote.grand_total else "0%"],
        ["3", "三类 规费 + 税金", _format_amount(quote.category3),
         f"{quote.category3 / quote.grand_total * 100:.2f}%" if quote.grand_total else "0%"],
        ["", "工程总造价", _format_amount(quote.grand_total), "100.00%"],
    ]
    _add_table(doc, ["序号", "费用类别", "金额(元)", "占比"],
               summary_rows, col_widths_cm=[2, 8, 4, 3])
    doc.add_page_break()

    # ------- 分部分项工程量清单计价表 -------
    _add_title(doc, "分部分项工程量清单计价表")
    cat1_rows = []
    for idx, comp in enumerate(quote.items, start=1):
        cat1_rows.append([
            idx, comp.item.item_name, comp.item.specialty or "",
            comp.item.unit, comp.item.qty, f"{comp.item.price:.2f}",
            _format_amount(comp.total), comp.item.remark or "",
        ])
    cat1_rows.append(["", "本页小计 / 一类合计", "", "", "", "",
                      _format_amount(quote.category1), ""])
    _add_table(doc,
        ["序号", "项目名称", "专业", "单位", "工程量", "综合单价(元)", "合价(元)", "备注"],
        cat1_rows, col_widths_cm=[1.5, 5, 2, 1.2, 1.5, 1.8, 2.5, 3])
    doc.add_page_break()

    # ------- 措施项目清单计价表 -------
    _add_title(doc, "措施项目清单计价表(二类)")
    cat2_rows = []
    for idx, (name, amount) in enumerate(quote.category2_detail.items(), start=1):
        rate = amount / quote.category1 * 100 if quote.category1 else 0
        cat2_rows.append([idx, name, "分部分项合计", f"{rate:.4f}%", _format_amount(amount), ""])
    cat2_rows.append(["", "二类合计", "", "", _format_amount(quote.category2), ""])
    _add_table(doc,
        ["序号", "措施项目名称", "计算基数", "费率(%)", "金额(元)", "备注"],
        cat2_rows, col_widths_cm=[1.5, 5, 3, 2, 3, 2.5])
    doc.add_page_break()

    # ------- 规费/税金项目计价表 -------
    _add_title(doc, "规费/税金项目计价表(三类)")
    cat3_rows = []
    for idx, (name, amount) in enumerate(quote.category3_detail.items(), start=1):
        if "规费" in name:
            base = quote.category1 + quote.category2
            rate = amount / base * 100 if base else 0
            base_desc = "分部分项+措施"
        else:
            base = quote.category1 + quote.category2 + quote.category3_detail.get("规费", 0)
            rate = amount / base * 100 if base else 0
            base_desc = "分部分项+措施+规费"
        cat3_rows.append([idx, name, base_desc, f"{rate:.4f}%", _format_amount(amount), ""])
    cat3_rows.append(["", "三类合计", "", "", _format_amount(quote.category3), ""])
    _add_table(doc,
        ["序号", "项目名称", "计算基数", "费率(%)", "金额(元)", "备注"],
        cat3_rows, col_widths_cm=[1.5, 5, 3, 2, 3, 2.5])

    doc.save(output_path)
    return str(Path(output_path).resolve())
